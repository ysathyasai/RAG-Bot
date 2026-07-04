"""
document_processor.py — Document loading, parsing, and chunking.
Supports PDF, DOCX, TXT, and plain text content.
"""

import os
import uuid
import logging
from typing import List, Tuple
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)

# ─── Chunking Settings ────────────────────────────────────────────────────────
CHUNK_SIZE = 600
CHUNK_OVERLAP = 80


class DocumentProcessor:
    """
    Loads and chunks documents from various file types.
    Returns LangChain Document objects with rich metadata.
    """

    def __init__(self):
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    def process_file(self, file_path: str, original_filename: str) -> Tuple[str, List[Document]]:
        """
        Process a file and return (doc_id, list_of_chunks).
        Each chunk is a LangChain Document with metadata.
        """
        ext = os.path.splitext(original_filename)[1].lower()
        doc_id = str(uuid.uuid4())

        logger.info(f"Processing file: {original_filename} (type={ext}, doc_id={doc_id})")

        if ext == ".pdf":
            raw_docs = self._load_pdf(file_path, original_filename)
        elif ext in (".docx", ".doc"):
            raw_docs = self._load_docx(file_path, original_filename)
        elif ext in (".txt", ".md", ".csv"):
            raw_docs = self._load_text(file_path, original_filename)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        chunks = self._chunk_documents(raw_docs, doc_id, original_filename, ext)
        logger.info(f"Created {len(chunks)} chunks from '{original_filename}'")
        return doc_id, chunks

    def process_text(self, text: str, source_name: str = "Pasted Text") -> Tuple[str, List[Document]]:
        """Process raw pasted text content."""
        doc_id = str(uuid.uuid4())
        raw_doc = Document(
            page_content=text,
            metadata={"source": source_name, "page": 1},
        )
        chunks = self._chunk_documents([raw_doc], doc_id, source_name, ".txt")
        return doc_id, chunks

    # ── Private Loaders ───────────────────────────────────────────────────────

    def _load_pdf(self, file_path: str, filename: str) -> List[Document]:
        """Load a PDF file page by page."""
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            docs = []
            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    docs.append(
                        Document(
                            page_content=text,
                            metadata={"source": filename, "page": page_num},
                        )
                    )
            return docs
        except Exception as e:
            logger.error(f"PDF loading failed: {e}")
            raise

    def _load_docx(self, file_path: str, filename: str) -> List[Document]:
        """Load a DOCX file, treating paragraphs as pages."""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            full_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return [Document(page_content=full_text, metadata={"source": filename, "page": 1})]
        except Exception as e:
            logger.error(f"DOCX loading failed: {e}")
            raise

    def _load_text(self, file_path: str, filename: str) -> List[Document]:
        """Load a plain text / markdown file."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
            return [Document(page_content=content, metadata={"source": filename, "page": 1})]
        except Exception as e:
            logger.error(f"Text loading failed: {e}")
            raise

    # ── Private Chunking ──────────────────────────────────────────────────────

    def _chunk_documents(
        self,
        docs: List[Document],
        doc_id: str,
        filename: str,
        file_type: str,
    ) -> List[Document]:
        """Split documents into chunks and enrich their metadata."""
        chunks = self.splitter.split_documents(docs)
        for i, chunk in enumerate(chunks):
            chunk.metadata.update(
                {
                    "doc_id": doc_id,
                    "filename": filename,
                    "file_type": file_type,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                }
            )
        return chunks
